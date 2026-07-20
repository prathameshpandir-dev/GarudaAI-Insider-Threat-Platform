import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding/margins for table cells in dxa (1/20th of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Configure Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Color Palette Definitions
    PRIMARY_COLOR = RGBColor(99, 106, 241)   # #6366F1 Indigo Accent
    DARK_TEXT = RGBColor(31, 41, 55)         # #1F2937 Off-Black
    GRAY_TEXT = RGBColor(75, 85, 99)         # #4B5563 Muted Gray
    
    # Configure Default Paragraph Style Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT

    # ==========================================
    # TITLE HEADER SECTION
    # ==========================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("GARUDAAI SECURITY PLATFORM")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("System Specifications & Product Architecture Guide")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = GRAY_TEXT
    
    doc.add_paragraph("_" * 60) # Divider Line
    doc.add_paragraph("")

    # ==========================================
    # 1. INTRODUCTION & PROBLEM SOLVED
    # ==========================================
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Introduction")
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph(
        "GarudaAI is an enterprise-grade insider threat detection and behavior trust score analytics platform. "
        "Unlike external security threats that target firewall boundaries, insider threats originate from compromised credentials "
        "or disgruntled internal personnel already trusted within the network. Standard SIEM systems log millions of actions, "
        "but fail to relate isolated logs into chronological narratives, causing severe alert fatigue."
    )
    
    doc.add_paragraph(
        "GarudaAI correlates isolated database reads, logins, USB file transfers, and web proxies to build a comprehensive "
        "Employee Behavior Trust Score. If score anomalies trigger alerts, the platform launches Google Gemini 1.5 Flash "
        "models to generate interactive incident summaries and automated security playbook recommendations."
    )

    # ==========================================
    # 2. KEY FEATURES
    # ==========================================
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Core Platform Features")
    h2_run.font.bold = True
    h2_run.font.color.rgb = PRIMARY_COLOR

    features_table = doc.add_table(rows=1, cols=3)
    features_table.autofit = False
    
    # Headers
    hdr_cells = features_table.rows[0].cells
    hdr_cells[0].text = "Feature Module"
    hdr_cells[1].text = "Technical Mechanism"
    hdr_cells[2].text = "Operational Benefit"
    
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, "6366F1")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    features_data = [
        ("Behavior Trust Engine", 
         "Delta subtraction framework combined with recovery algorithms that add +1.0 trust score point for each calendar day of clean network actions.", 
         "Establishes a quantitative risk level for every user profile, shifting security auditing from passive logs to active risks."),
        
        ("Timeline Collapsing", 
         "Linear grouping module that condenses routine logins and browser pings, while keeping high-risk activities expanded and color-coded.", 
         "Saves analysts hours of timeline inspection by filtering out 98% of standard network activity noise."),
        
        ("AI Playbook Assistant", 
         "Context prompt template loader feeding security histories to Google Gemini Flash models to return markdown-formatted mitigation actions.", 
         "Delivers structured response instructions instantly to SOC teams, enabling immediate isolation of compromised nodes."),
        
        ("Attack Simulator", 
         "In-memory scenario runner that injects mock exfiltration events and automatically recalculates scores and alerts.", 
         "Allows operators to stress-test threat thresholds and configure alert integrations without risking real network data.")
    ]
    
    for feat, mech, ben in features_data:
        row_cells = features_table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = mech
        row_cells[2].text = ben
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_background(cell, "F9FAFB") # Soft gray zebra stripe

    doc.add_paragraph("")

    # ==========================================
    # 3. TECHNICAL ARCHITECTURE & STACK
    # ==========================================
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. System Architecture & Tech Stack")
    h3_run.font.bold = True
    h3_run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph(
        "GarudaAI is built as a single-page decoupled web application. The frontend communicates with the backend Flask API "
        "using standard HTTP REST routing. If the MongoDB service is unreachable during local evaluation, a transparent "
        "database connector falls back to file-based JSON caching without interrupting service routes."
    )
    
    # Tech Stack Table
    tech_table = doc.add_table(rows=1, cols=2)
    tech_hdr = tech_table.rows[0].cells
    tech_hdr[0].text = "Component Layer"
    tech_hdr[1].text = "Technologies Selected & Role"
    
    for cell in tech_hdr:
        set_cell_background(cell, "374151")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    tech_data = [
        ("Frontend Client", "React 19, Vite, Tailwind CSS v3, Outfit Typography, FontAwesome Icons"),
        ("Data Visualization", "Chart.js, react-chartjs-2 (Dynamic responsive line trends)"),
        ("Backend Services", "Python, Flask, Flask-CORS (Cross-Origin Resource Sharing)"),
        ("Rate Limiter", "Flask-Limiter (Rate limits critical POST routes to 10 req/min)"),
        ("Primary Database", "MongoDB Community Server / Fallback JSON Caches"),
        ("AI Prompt Engine", "Google Gemini 1.5 Flash API (google-generativeai client sdk)"),
        ("User Authentication", "Firebase Admin SDK middleware tokens validation"),
        ("Testing Suite", "Python unittest (Automated endpoint & scoring tests)")
    ]
    
    for layer, tech in tech_data:
        row = tech_table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        for cell in row:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_background(cell, "F3F4F6")
            
    doc.add_paragraph("")

    # ==========================================
    # 4. DATABASE & COLLECTION SCHEMA
    # ==========================================
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Database & Collection Schema")
    h4_run.font.bold = True
    h4_run.font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph(
        "The system maps operational metrics to three core collections. Indexes are built on unique IDs and "
        "foreign employee IDs to maintain rapid query execution during timeline correlation searches."
    )

    doc.add_heading("4.1 Employees Collection", level=2)
    doc.add_paragraph(
        "Stores the baseline metadata and active Behavior Trust Score for every user:\n"
        "• employee_id (String, Indexed, Unique Key)\n"
        "• full_name (String)\n"
        "• department (String)\n"
        "• role (String)\n"
        "• is_privileged_user (Boolean)\n"
        "• current_score (Float, Default: 100.0)"
    )

    doc.add_heading("4.2 Events Collection", level=2)
    doc.add_paragraph(
        "Chronological activity log registry:\n"
        "• event_id (String, Indexed, Unique Key)\n"
        "• employee_id (String, Indexed)\n"
        "• timestamp (Datetime)\n"
        "• type (String: logon, file, device, http, email, privilege)\n"
        "• details (Nested Document: location, file_size_mb, access_level, etc.)"
    )

    doc.add_heading("4.3 Alerts Collection", level=2)
    doc.add_paragraph(
        "Incidents generated by the behavior engine or simulation injections:\n"
        "• alert_id (String, Indexed, Unique Key)\n"
        "• employee_id (String, Indexed)\n"
        "• timestamp (Datetime)\n"
        "• type (String: USB Theft, Mass File Download, Impossible Travel)\n"
        "• severity (String: Critical, High, Medium, Low)\n"
        "• description (String)\n"
        "• status (String: Open, Closed)\n"
        "• ai_explanation (Markdown String, cached response)"
    )
    
    doc.add_paragraph("")

    # ==========================================
    # 5. API SPECIFICATIONS
    # ==========================================
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. REST API Specifications")
    h5_run.font.bold = True
    h5_run.font.color.rgb = PRIMARY_COLOR
    
    api_table = doc.add_table(rows=1, cols=4)
    api_hdr = api_table.rows[0].cells
    api_hdr[0].text = "Method"
    api_hdr[1].text = "Endpoint Route"
    api_hdr[2].text = "Description"
    api_hdr[3].text = "Auth Type"
    
    for cell in api_hdr:
        set_cell_background(cell, "6366F1")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    api_data = [
        ("GET", "/api/health", "Retrieves system health and DB connection state.", "None"),
        ("GET", "/api/employees", "Returns employee profiles sorted by risk level.", "Bearer JWT / Dev Bypass"),
        ("GET", "/api/employees/<id>/timeline", "Returns parsed and collapsed user log list.", "Bearer JWT / Dev Bypass"),
        ("GET", "/api/employees/<id>/trust-score/history", "Returns historical snapshots for line charts.", "Bearer JWT / Dev Bypass"),
        ("GET", "/api/alerts", "Retrieves active threat alert logs.", "Bearer JWT / Dev Bypass"),
        ("POST", "/api/simulate", "Injects simulated incident attack scenarios.", "Bearer JWT / Dev Bypass"),
        ("POST", "/api/chat", "Translates natural language questions to DB queries.", "Bearer JWT / Dev Bypass")
    ]
    
    for method, route, desc, auth in api_data:
        row = api_table.add_row().cells
        row[0].text = method
        row[1].text = route
        row[2].text = desc
        row[3].text = auth
        for cell in row:
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_background(cell, "F9FAFB")
            
    doc.add_paragraph("")

    # ==========================================
    # 6. TESTING & ROADMAP
    # ==========================================
    h6 = doc.add_heading(level=1)
    h6_run = h6.add_run("6. Quality Assurance & Future Roadmap")
    h6_run.font.bold = True
    h6_run.font.color.rgb = PRIMARY_COLOR

    doc.add_heading("6.1 Testing Summary", level=2)
    doc.add_paragraph(
        "The system has been validated by an automated backend test harness covering:\n"
        "• Behavior trust scoring calculations (e.g. correct delta points for after-hours authentication).\n"
        "• Integration test routes verifying HTTP 200 checks across employee lists and timelines.\n"
        "• End-to-end simulation flow (reset database, inject mass download logs, assert alert publishing, and check score updates)."
    )

    doc.add_heading("6.2 Future Enhancements (v2 Roadmap)", level=2)
    doc.add_paragraph(
        "1. Active Directory Integration: Hook the playbook generator to AD LDAP APIs to lock user sessions automatically.\n"
        "2. Stream Processing Ingestion: Integrate with Apache Kafka or Splunk for real-time streaming SIEM events parsing.\n"
        "3. Anomaly Detection ML: Add Isolation Forest or Autoencoder models to learn custom, employee-specific baselines."
    )

    # Save Document
    filename = "GarudaAI_Project_Documentation.docx"
    doc.save(filename)
    print(f"Native Word Document saved as: {filename}")

if __name__ == "__main__":
    create_document()
